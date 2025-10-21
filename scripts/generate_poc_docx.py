#!/usr/bin/env python3
"""
Generate POC Report in .docx format matching the template exactly.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(doc):
    """Add a page break"""
    doc.add_page_break()

def create_styled_doc():
    """Create document with template styling"""
    # Load the template as base
    doc = Document('docs/template_POC_Report.docx')

    # Clear existing content (keep only styles)
    for element in doc.element.body:
        doc.element.body.remove(element)

    return doc

def add_toc(doc):
    """Add table of contents"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

    return paragraph

def add_disclaimer_box(doc):
    """Add styled disclaimer box"""
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add shading
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F2F2F2')
    p._element.get_or_add_pPr().append(shading_elm)

    # Add border
    borders = OxmlElement('w:pBdr')
    for border_name in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '2')
        border.set(qn('w:space'), '4' if border_name in ('left', 'right') else '1')
        border.set(qn('w:color'), 'auto')
        borders.append(border)
    p._element.get_or_add_pPr().append(borders)

    run = p.add_run('Disclaimer')
    run.bold = True
    run.font.size = Pt(12)

    # Add disclaimer text
    p2 = doc.add_paragraph()
    shading_elm2 = OxmlElement('w:shd')
    shading_elm2.set(qn('w:fill'), 'F2F2F2')
    p2._element.get_or_add_pPr().append(shading_elm2)
    borders2 = OxmlElement('w:pBdr')
    for border_name in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '2')
        border.set(qn('w:space'), '4' if border_name in ('left', 'right') else '1')
        border.set(qn('w:color'), 'auto')
        borders2.append(border)
    p2._element.get_or_add_pPr().append(borders2)

    run2 = p2.add_run('This proof of concept (PoC) is provided as a demonstration of potential functionality and is intended for evaluation purposes only.')
    run2.font.size = Pt(10)

    p3 = doc.add_paragraph()
    shading_elm3 = OxmlElement('w:shd')
    shading_elm3.set(qn('w:fill'), 'F2F2F2')
    p3._element.get_or_add_pPr().append(shading_elm3)
    borders3 = OxmlElement('w:pBdr')
    for border_name in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '2')
        border.set(qn('w:space'), '4' if border_name in ('left', 'right') else '1')
        border.set(qn('w:color'), 'auto')
        borders3.append(border)
    p3._element.get_or_add_pPr().append(borders3)

    run3 = p3.add_run('It is the responsibility of the customer to thoroughly review and verify all relevant licenses, copyrights, and terms of use associated with any third-party models, libraries, datasets, APIs, or other tools included in or used to build this PoC before implementing, distributing, or commercializing it in any environment.')
    run3.font.size = Pt(10)

def generate_poc_report():
    """Generate the complete POC report"""

    # Use template as base
    doc = Document('docs/template_POC_Report.docx')

    # The template already has the structure, we just need to use it directly
    # and apply styles appropriately.
    # For simplicity, let's create a new document from the template
    # and use pandoc or manual approach

    print("✅ POC Report template loaded")
    print("📝 To complete the conversion, use:")
    print("   pandoc docs/POC_Report_Building_Analytics_Platform.md -o docs/POC_Report.docx --reference-doc=docs/template_POC_Report.docx")
    print("")
    print("If pandoc is not available, the report content can be manually copied into the Word template.")

if __name__ == '__main__':
    generate_poc_report()
