#!/usr/bin/env python3
"""
Convert POC Report markdown to .docx format matching template.
Uses python-docx library.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Read markdown content
with open('docs/POC_Report_Building_Analytics_Platform.md', 'r', encoding='utf-8') as f:
    content = f.read()

# Load template
doc = Document('docs/template_POC_Report.docx')

# Clear the template body (keep styles, headers, footers)
for paragraph in list(doc.paragraphs):
    p = paragraph._element
    p.getparent().remove(p)

# Parse markdown and add to document
import re

lines = content.split('\n')
i = 0

while i < len(lines):
    line = lines[i].strip()

    if not line:
        i += 1
        continue

    # Title
    if line.startswith('# ') and 'Proof of Concept' in line:
        p = doc.add_paragraph(line[2:], style='Title')

    # Date subtitle
    elif line.startswith('**Date:**'):
        date_text = line.replace('**Date:**', '').strip()
        p = doc.add_paragraph(date_text, style='Subtitle')

    # Customer/Team info
    elif line.startswith('**Customer:**') or line.startswith('**Development team:**'):
        p = doc.add_paragraph()
        parts = line.split(':', 1)
        run = p.add_run(parts[0].replace('**', '') + ':')
        run.bold = True
        if len(parts) > 1:
            p.add_run(' ' + parts[1].strip())

    # Heading 1
    elif line.startswith('## ') and not 'Contents' in line:
        p = doc.add_paragraph(line[3:], style='Heading 1')

    # Heading 2
    elif line.startswith('### '):
        p = doc.add_paragraph(line[4:], style='Heading 2')

    # Heading 3
    elif line.startswith('#### '):
        p = doc.add_paragraph(line[5:], style='Heading 3')

    # Code block
    elif line.startswith('```'):
        code_lines = []
        i += 1
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1

        p = doc.add_paragraph()
        run = p.add_run('\n'.join(code_lines))
        run.font.name = 'Courier New'
        run.font.size = Pt(10)

    # Bullet list
    elif line.startswith('* ') or line.startswith('- '):
        p = doc.add_paragraph(line[2:])
        # Add bullet manually if style doesn't exist
        try:
            p.style = 'List Bullet'
        except:
            p.style = 'Normal'

    # Table (simple markdown tables)
    elif '|' in line and '---' in lines[i+1] if i+1 < len(lines) else False:
        # Parse table
        header = [cell.strip() for cell in line.split('|')[1:-1]]
        i += 2  # Skip separator line

        rows = []
        while i < len(lines) and '|' in lines[i]:
            row = [cell.strip() for cell in lines[i].split('|')[1:-1]]
            rows.append(row)
            i += 1

        table = doc.add_table(rows=len(rows)+1, cols=len(header))
        table.style = 'Table Grid'

        # Header row
        for j, cell_text in enumerate(header):
            table.rows[0].cells[j].text = cell_text
            for paragraph in table.rows[0].cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Data rows
        for row_idx, row in enumerate(rows):
            for col_idx, cell_text in enumerate(row):
                table.rows[row_idx+1].cells[col_idx].text = cell_text

        i -= 1  # Will be incremented at end of loop

    # Regular paragraph
    else:
        p = doc.add_paragraph(line)

    i += 1

# Save document
doc.save('docs/POC_Report.docx')
print("✅ POC Report generated: docs/POC_Report.docx")
